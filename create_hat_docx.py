import os
import requests
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Download a golden yellow 3D graduation cap icon from Icons8
url = "https://img.icons8.com/color/480/000000/graduation-cap.png"
hat_path = r"c:\Users\asus\budgetbeam\topscore_hat.png"
try:
    res = requests.get(url)
    with open(hat_path, "wb") as f:
        f.write(res.content)
except Exception as e:
    print("Download failed", e)

# Regenerate Docx
doc = docx.Document()
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_centered(text, size=16, bold=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold

add_centered("DEPARTMENT OF COMPUTER SCIENCE", 16, False, 6)
add_centered("RAJAGIRI COLLEGE OF SOCIAL SCIENCES", 16, False, 6)
add_centered("(AUTONOMOUS)", 16, False, 6)
add_centered("KALAMASSERY-KOCHI-683104", 14, False, 24)

p_img1 = doc.add_paragraph()
p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
rcss_path = r"c:\Users\asus\budgetbeam\RCSS.png"
if os.path.exists(rcss_path):
    run_img1 = p_img1.add_run()
    run_img1.add_picture(rcss_path, width=Inches(1.3))
p_img1.paragraph_format.space_after = Pt(24)

add_centered("BSc Computer Science (Hons.) Data Analytics", 16, False, 18)
add_centered("WPL Project Report", 16, False, 18)
add_centered("TopScore", 26, True, 6)
add_centered("Quiz Management System", 16, False, 30)

table = doc.add_table(rows=3, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

details = [
    ("Name", ":", "Amal Biju"),
    ("Semester", ":", "4th Semester"),
    ("Register No.", ":", "24118005")
]

for i, (col1, col2, col3) in enumerate(details):
    row_cells = table.rows[i].cells
    p1 = row_cells[0].paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run(col1); r1.font.name = 'Calibri'; r1.font.size = Pt(14); r1.bold = True
    p2 = row_cells[1].paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(col2); r2.font.name = 'Calibri'; r2.font.size = Pt(14); r2.bold = True
    p3 = row_cells[2].paragraphs[0]; p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r3 = p3.add_run(col3); r3.font.name = 'Calibri'; r3.font.size = Pt(14)

for cell in table.columns[0].cells: cell.width = Inches(1.2)
for cell in table.columns[1].cells: cell.width = Inches(0.3)
for cell in table.columns[2].cells: cell.width = Inches(2.0)

p_space = doc.add_paragraph()
p_space.paragraph_format.space_before = Pt(30)

if os.path.exists(hat_path):
    # Ensure it's a true PNG just in case
    try:
        from PIL import Image
        img = Image.open(hat_path).convert("RGBA")
        true_hat = r"c:\Users\asus\budgetbeam\topscore_hat_true.png"
        img.save(true_hat, "PNG")
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img2 = p_img2.add_run()
        run_img2.add_picture(true_hat, width=Inches(4.5))
    except Exception as e:
        print("PIL error processing hat:", e)

doc.save(r"c:\Users\asus\budgetbeam\TopScore_Cover_Page.docx")
print("Docx regeneration complete!")

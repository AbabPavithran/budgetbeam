import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = docx.Document()

# Adjust margins to fit everything nicely on a cover page
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_centered(text, size=16, bold=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold

add_centered("DEPARTMENT OF COMPUTER SCIENCE", 16, False, 6)
add_centered("RAJAGIRI COLLEGE OF SOCIAL SCIENCES", 16, False, 6)
add_centered("(AUTONOMOUS)", 16, False, 6)
add_centered("KALAMASSERY-KOCHI-683104", 14, False, 24)

# Internal Logo
p_img1 = doc.add_paragraph()
p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
rcss_path = r"c:\Users\asus\budgetbeam\Budgetbeam_Overleaf_Ready__2__Extracted\RCSS.png"
if os.path.exists(rcss_path):
    run_img1 = p_img1.add_run()
    run_img1.add_picture(rcss_path, width=Inches(1.3))
p_img1.paragraph_format.space_after = Pt(24)

add_centered("BSc Computer Science (Hons.) Data Analytics", 16, False, 18)
add_centered("WPL Project Report", 16, False, 18)
add_centered("<project name>", 16, False, 40)

# Grid setup for details
table = doc.add_table(rows=3, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

details = [
    ("Name", ":", "Abab P K"),
    ("Semester", ":", "<sem>"),
    ("Register No.", ":", "24118002")
]

for i, (col1, col2, col3) in enumerate(details):
    row_cells = table.rows[i].cells
    
    # Left column Setup
    p1 = row_cells[0].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = p1.add_run(col1)
    run1.font.name = 'Calibri'
    run1.font.size = Pt(14)
    
    # Center colon Setup
    p2 = row_cells[1].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(col2)
    run2.font.name = 'Calibri'
    run2.font.size = Pt(14)
    
    # Right column Setup
    p3 = row_cells[2].paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run3 = p3.add_run(col3)
    run3.font.name = 'Calibri'
    run3.font.size = Pt(14)

for cell in table.columns[0].cells: cell.width = Inches(1.2)
for cell in table.columns[1].cells: cell.width = Inches(0.3)
for cell in table.columns[2].cells: cell.width = Inches(2.0)

# Add Spacing before the final customized image
p_space = doc.add_paragraph()
p_space.paragraph_format.space_before = Pt(30)

# Large custom image at bottom
p_img2 = doc.add_paragraph()
p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
custom_img = r"c:\Users\asus\budgetbeam\budgetbeam_custom_logo_word.png"
if os.path.exists(custom_img):
    try:
        from PIL import Image
        img = Image.open(custom_img).convert("RGBA")
        true_png = r"c:\Users\asus\budgetbeam\budgetbeam_true.png"
        img.save(true_png, "PNG")
        run_img2 = p_img2.add_run()
        run_img2.add_picture(true_png, width=Inches(5.0))
    except Exception as e:
        print("Error processing image:", e)

doc.save(r"c:\Users\asus\budgetbeam\Budgetbeam_Cover_Page.docx")
